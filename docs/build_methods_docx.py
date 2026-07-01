"""Render docs/01_DrStone_Methods.md to a .docx with embedded figures.

Handles the subset of Markdown this writeup uses: #/##/### headings, **bold**
inline runs, > blockquotes, and [[FIGURE: name.png]] placeholders (embedded from
docs/figures/). Run:  python docs/build_methods_docx.py
"""

from __future__ import annotations

import os
import re

from docx import Document
from docx.shared import Inches, Pt

import sys

HERE = os.path.dirname(os.path.abspath(__file__))
_STEM = sys.argv[1] if len(sys.argv) > 1 else "01_DrStone_Methods"
MD = os.path.join(HERE, _STEM + ".md")
DOCX = os.path.join(HERE, _STEM + ".docx")
FIGS = os.path.join(HERE, "figures")

_MDTABLE = re.compile(r"^\s*\|.*\|\s*$")

_BOLD = re.compile(r"\*\*(.+?)\*\*")
_FIG = re.compile(r"\[\[FIGURE:\s*(.+?)\s*\]\]")


def add_runs(par, text):
    """Add text to a paragraph, rendering **bold** spans as bold runs."""
    pos = 0
    for m in _BOLD.finditer(text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        par.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])


def main():
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)
    with open(MD, encoding="utf-8") as f:
        lines = f.read().splitlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        i += 1
        if not line.strip():
            continue
        # ---- markdown table block ----
        if _MDTABLE.match(line):
            block = [line]
            while i < len(lines) and _MDTABLE.match(lines[i]):
                block.append(lines[i].rstrip()); i += 1

            def cells(row):
                return [c.strip() for c in row.strip().strip("|").split("|")]
            rows = [cells(r) for r in block
                    if not set(r.replace("|", "").strip()) <= set("-: ")]
            if rows:
                t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                t.style = "Light Grid Accent 1"
                for ri, r in enumerate(rows):
                    for ci, txt in enumerate(r[:len(rows[0])]):
                        cell = t.rows[ri].cells[ci]
                        cell.paragraphs[0].text = ""
                        run = cell.paragraphs[0].add_run(_BOLD.sub(r"\1", txt))
                        if ri == 0 or txt.startswith("**"):
                            run.bold = True
                        run.font.size = Pt(8)
            continue
        fig = _FIG.match(line.strip())
        if fig:
            path = os.path.join(FIGS, fig.group(1))
            if os.path.exists(path):
                doc.add_picture(path, width=Inches(6.0))
            else:
                doc.add_paragraph(f"[missing figure: {fig.group(1)}]")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, line[2:])
        else:
            add_runs(doc.add_paragraph(), line)

    doc.save(DOCX)
    n_imgs = sum(1 for _ in doc.inline_shapes)
    print(f"wrote {DOCX}  ({len(doc.paragraphs)} paragraphs, {n_imgs} figures)")


if __name__ == "__main__":
    main()
